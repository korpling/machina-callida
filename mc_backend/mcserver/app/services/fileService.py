import os
from io import StringIO
from random import shuffle
from tempfile import mkstemp
from typing import List, Set
import rapidjson as json
import conllu
from bs4 import BeautifulSoup
from conllu import TokenList
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from xhtml2pdf import pisa

from mcserver import Config
from mcserver.app.models import DownloadableFile, FileType, Solution, Exercise, ExerciseType, SolutionElement, \
    VocabularyCorpus
from mcserver.app.services import TextService, XMLservice


class FileService:
    """Service for loading content from files and serving it as downloadable data."""

    downloadable_files: List[DownloadableFile] = []

    @staticmethod
    def create_tmp_file(file_type: FileType, file_id: str) -> DownloadableFile:
        """ Creates a new temporary file and adds it to the FileService watch list. """
        # generate temp file
        (handle, path) = mkstemp(suffix=".{0}".format(file_type.value), dir=Config.TMP_DIRECTORY)
        # grant all permissions for the file to everybody, so Docker can handle the files during builds / updates
        os.fchmod(handle, 0o777)
        file_name: str = os.path.basename(path)
        existing_file: DownloadableFile = DownloadableFile(file_id=file_id, file_name=file_name, file_type=file_type,
                                                           file_path=path)
        FileService.downloadable_files.append(existing_file)
        return existing_file

    @staticmethod
    def get_body_for_matching_exercise(solutions_shuffled: List[Solution]) -> str:
        """Builds an HTML table for a matching exercise, with keys in one column and values in the other."""
        table_string: str = "<table>{0}</table>"
        row_string_template: str = "<tr><td>{0}</td><td>{1}</td></tr>"
        row_strings: List[str] = []
        for solution in solutions_shuffled:
            row_strings.append(row_string_template.format(solution.target.content, solution.value.content))
        return table_string.format("".join(row_strings))

    @staticmethod
    def get_file_content(path: str) -> str:
        """Open a file and return its content as a string."""
        ret_val = ""
        with open(path, "r", encoding="utf-8") as f:
            ret_val = f.read()
        return ret_val

    @staticmethod
    def get_pdf_html_string(exercise: Exercise, conll: List[TokenList], file_type: FileType,
                            solutions: List[Solution]) -> str:
        """Builds a HTML string for exercises to be exported to PDF."""
        solutions_shuffled: List[Solution] = FileService.shuffle_solutions(solutions)
        if exercise.exercise_type == ExerciseType.cloze.value:
            text_with_gaps = XMLservice.get_moodle_cloze_text_with_solutions(conll, file_type, solutions)
            solution_string: str = " ".join([solution.target.content for solution in solutions_shuffled])
            return f"<p>{exercise.exercise_type_translation}: {exercise.instructions}</p><p>{text_with_gaps}</p><br><br><div>{solution_string}</div>"
        elif exercise.exercise_type == ExerciseType.matching.value:
            exercise_body: str = FileService.get_body_for_matching_exercise(solutions_shuffled)
            return f"<br><p>{exercise.exercise_type_translation}: {exercise.instructions}</p><p>{exercise_body}</p>"
        elif exercise.exercise_type == ExerciseType.markWords.value:
            text_with_gaps: str = TextService.strip_whitespace(" ".join([y["form"] for x in conll for y in x]))
            return f"<p>{exercise.exercise_type_translation}: {exercise.instructions}</p><p>{text_with_gaps}</p><br><br>"

    @staticmethod
    def get_reff_from_disk(disk_urn: str) -> List[str]:
        """ Loads valid references for a standard CTS corpus from disk. """
        file_path: str = os.path.join(Config.REFF_CACHE_DIRECTORY, disk_urn)
        reff: List[str] = []
        if os.path.isfile(file_path):
            reff_json: str = FileService.get_file_content(file_path)
            reff = json.loads(reff_json)
        return reff

    @staticmethod
    def get_vocabulary_set(vocabulary_corpus: VocabularyCorpus, frequency_upper_bound: int) -> Set[str]:
        """ Retrieves a set of highly frequent lemmata from a specific vocabulary. """
        vocabulary_file_content: str = FileService.get_file_content(
            os.path.join(Config.ASSETS_DIRECTORY, vocabulary_corpus.value))
        vocabulary_list: List[str] = json.loads(vocabulary_file_content)
        return set(vocabulary_list[:frequency_upper_bound])

    @staticmethod
    def make_docx_file(exercise: Exercise, path: str, conll: List[TokenList], file_type: FileType,
                       solutions: List[Solution]) -> None:
        """ Creates a .docx file for a given exercise. """
        doc: Document = Document()
        par1: Paragraph = doc.add_paragraph(exercise.exercise_type_translation)
        par2: Paragraph = doc.add_paragraph(exercise.instructions)
        if exercise.exercise_type == ExerciseType.cloze.value:
            text_with_gaps = XMLservice.get_moodle_cloze_text_with_solutions(conll, file_type, solutions)
            par3: Paragraph = doc.add_paragraph(text_with_gaps)
            par4: Paragraph = doc.add_paragraph(
                " ".join([x.target.content for x in FileService.shuffle_solutions(solutions)]))
        elif exercise.exercise_type == ExerciseType.matching.value:
            solutions_shuffled: List[Solution] = FileService.shuffle_solutions(solutions)
            table1 = doc.add_table(rows=len(solutions_shuffled), cols=2)
            for i in range(len(solutions_shuffled)):
                table1.rows[i].cells[0].text = solutions_shuffled[i].target.content
                table1.rows[i].cells[1].text = solutions_shuffled[i].value.content
        elif exercise.exercise_type == ExerciseType.markWords.value:
            text_with_gaps: str = TextService.strip_whitespace(" ".join([y["form"] for x in conll for y in x]))
            par3: Paragraph = doc.add_paragraph(text_with_gaps)
        doc.save(path)

    @staticmethod
    def make_tmp_file_from_exercise(file_type: FileType, exercise: Exercise,
                                    solution_indices: List[int] = None) -> DownloadableFile:
        """Creates a temporary file for the exercise data, so the users can download it."""
        existing_file: DownloadableFile = FileService.create_tmp_file(file_type, exercise.eid)
        conll: List[TokenList] = conllu.parse(exercise.conll)
        solutions: List[Solution] = [Solution(json_dict=x) for x in json.loads(exercise.solutions)]
        if solution_indices is not None:
            solutions = [solutions[x] for x in solution_indices]
        # write the relevant content to the file
        if file_type == FileType.pdf:
            html_string: str = FileService.get_pdf_html_string(exercise, conll, file_type, solutions)
            with open(existing_file.file_path, "wb+") as f:
                pdf = pisa.CreatePDF(StringIO(html_string), f)
        elif file_type == FileType.xml:
            # export exercise data to XML
            xml_string: str = XMLservice.create_xml_string(exercise, conll, file_type, solutions)
            with open(existing_file.file_path, "w+") as f:
                f.write(xml_string)
        elif file_type == FileType.docx:
            FileService.make_docx_file(exercise, existing_file.file_path, conll, file_type, solutions)
        return existing_file

    @staticmethod
    def make_tmp_file_from_html(urn: str, file_type: FileType, html_content: str) -> DownloadableFile:
        """Creates a temporary file for the HTML content, so the users can download it."""
        existing_file: DownloadableFile = FileService.create_tmp_file(file_type, urn)
        # write the relevant content to the file
        if file_type == FileType.pdf:
            with open(existing_file.file_path, "wb+") as f:
                pdf = pisa.CreatePDF(StringIO(html_content), f)
        elif file_type == FileType.docx:
            soup: BeautifulSoup = BeautifulSoup(html_content, 'html.parser')
            doc: Document = Document()
            par1: Paragraph = doc.add_paragraph(soup.p.text)
            par2: Paragraph = doc.add_paragraph()
            for span in soup.find_all("span", {"class": "tok"}):
                run: Run = par2.add_run(span.text.strip())
                if span.find("u") is not None:
                    run.underline = True
                par2.add_run(" ")
            doc.save(existing_file.file_path)
        return existing_file

    @staticmethod
    def shuffle_solutions(solutions: List[Solution]) -> List[Solution]:
        """Shuffles the key-value associations in the solution dictionary."""
        targets: List[SolutionElement] = [x.target for x in solutions]
        shuffle(targets)
        solutions_shuffled: List[Solution] = [Solution(target=targets[i], value=solutions[i].value) for i in
                                              range(len(targets))]
        return solutions_shuffled
